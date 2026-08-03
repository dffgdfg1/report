# -*- coding: utf-8 -*-
"""生成 骨架.docx：删除两个示例测试段落，保留固定骨架（封面→汇总→定义）+ 报告结束标记。
留下的正文块：0..67(定义结尾空段) + 报告结束标记 + 其后 + sectPr。
删除范围：从 '高温耐久测试' 段落 到 '报告结束' 段落之前。"""
import copy
from docx import Document

W='{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
SRC=r"C:\Users\admin\Desktop\报告生成器\模板库\主体_full.docx"
DST=r"C:\Users\admin\Desktop\报告生成器\模板库\骨架.docx"

doc=Document(SRC)
body=doc.element.body
children=list(body.iterchildren())

def ptext(el):
    if el.tag!=W+'p': return None
    return "".join([(x.text or "") for x in el.findall('.//'+W+'t')]).strip()

start=None; end=None
for idx,el in enumerate(children):
    t=ptext(el)
    if t=="高温耐久测试" and start is None:
        start=idx
    if t and "报告结束" in t:
        end=idx
        break
print("删除区间 blocks:", start, "..", end-1, " 报告结束在:", end)
assert start is not None and end is not None and start<end

for el in children[start:end]:
    body.remove(el)

doc.save(DST)

# 校验
d2=Document(DST)
print("骨架剩余表格数:", len(d2.tables))
for i,c in enumerate(d2.element.body.iterchildren()):
    if c.tag==W+'p':
        t="".join([(x.text or "") for x in c.findall('.//'+W+'t')]).strip()
        if t and ("测试" in t or "报告结束" in t or "汇总" in t):
            print("  P:", t[:30])
print("saved", DST)
