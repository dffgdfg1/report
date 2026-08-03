# -*- coding: utf-8 -*-
import io
from docx import Document
doc=Document(r"C:\Users\admin\Desktop\报告生成器\输出\_测试输出.docx")
out=io.StringIO()
# 测试段落表格从 TBL#3 起（前3张=封面/样品/汇总）
for ti in range(3,len(doc.tables)):
    t=doc.tables[ti]
    A_BLIP='{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
    nimg=len(t._tbl.findall('.//'+A_BLIP))
    out.write(f"\n--- TBL#{ti} {len(t.rows)}行 imgs={nimg} ---\n")
    for r in t.rows:
        seen=set(); cells=[]
        for c in r.cells:
            if id(c._tc) in seen: cells.append("<m>"); continue
            seen.add(id(c._tc))
            if c._tc.findall('.//'+A_BLIP): cells.append("[图]")
            else: cells.append(c.text.strip().replace("\n","/")[:18])
        out.write("  "+" | ".join(cells)+"\n")
open(r"C:\Users\admin\Desktop\报告生成器\构建\验证结果2.txt","w",encoding="utf-8").write(out.getvalue())
print("done")
