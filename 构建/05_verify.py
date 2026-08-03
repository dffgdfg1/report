# -*- coding: utf-8 -*-
import io
from docx import Document
from docx.shared import Emu
W='{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
A_EXT='{http://schemas.openxmlformats.org/drawingml/2006/main}ext'
A_BLIP='{http://schemas.openxmlformats.org/drawingml/2006/main}blip'

doc=Document(r"C:\Users\admin\Desktop\报告生成器\输出\_测试输出.docx")
out=io.StringIO()

# 图片尺寸统计
sizes={}
total=0
for ext in doc.element.body.findall('.//'+A_EXT):
    cx,cy=ext.get('cx'),ext.get('cy')
    if cx and cy:
        key=f"{Emu(int(cx)).cm:.2f} x {Emu(int(cy)).cm:.2f} cm"
        sizes[key]=sizes.get(key,0)+1
        total+=1
out.write(f"图片总数(按ext): {total}\n尺寸分布:\n")
for k,v in sizes.items():
    out.write(f"  {k}  ×{v}\n")

out.write(f"\n表格数: {len(doc.tables)}\n")
# 块序列关键标题
out.write("\n关键块:\n")
for c in doc.element.body.iterchildren():
    if c.tag==W+'p':
        t="".join([(x.text or '') for x in c.findall('.//'+W+'t')]).strip()
        if t and ('测试' in t or '报告结束' in t or '汇总' in t or '图片' in t or '报告编号' in t):
            out.write("  P: "+t[:36]+"\n")

# 汇总表
out.write("\n结果汇总表:\n")
tbl=doc.tables[2]
for r in tbl.rows:
    out.write("  "+" | ".join(c.text.strip().replace('\n','/')[:14] for c in r.cells)+"\n")

open(r"C:\Users\admin\Desktop\报告生成器\构建\验证结果.txt","w",encoding="utf-8").write(out.getvalue())
print("done")
