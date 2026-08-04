# -*- coding: utf-8 -*-
"""报告生成引擎：从骨架.docx + 格式供体(主体_full.docx) 生成完整报告。
- 填充封面/样品信息/结果汇总
- 逐个追加测试项目段落（克隆供体表格，保持官方排版）
- 试验图片表重建：2列，每张填满列宽(约7.6cm)、按真实比例、行间留白，每页约6张
"""
import os
import copy
from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn

# 图片布局：2列，每张填满列宽(约7.6cm)，高度按真实比例走(不拉伸变形)。
# 可用宽16cm / 2列 ≈ 8cm，图宽取7.6cm留少量边距；竖图高度封顶避免过高。
IMG_W_CM = 7.6        # 每张图占满列宽
IMG_MAX_H_CM = 6.5    # 高度上限（防竖图过高）
ROW_H_CM = 7.3        # 图片行高：留白使一页约放6张（3行×2列）
FIT_MODE = "fill_width"
# 兼容旧引用（/api/meta 用来显示角标）
IMG_H_CM = 5.7

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SKELETON = os.path.join(BASE, "模板库", "骨架.docx")
DONOR = os.path.join(BASE, "模板库", "主体_full.docx")

# ---------- 低层 XML 助手 ----------
def clone(el):
    return copy.deepcopy(el)

def para_text(p):
    return "".join([(t.text or "") for t in p.findall('.//' + W + 't')]).strip()

def iter_body(doc):
    return list(doc.element.body.iterchildren())

def find_para(doc, contains):
    """返回正文中第一个文本包含 contains 的段落元素。"""
    for el in doc.element.body.iterchildren():
        if el.tag == W + 'p' and contains in para_text(el):
            return el
    return None
# ---------- 单元格 / 文本 ----------
def cell_set_text(cell, text):
    set_tc_text(cell._tc, text)

def _set_para_align(p, align):
    """设置段落对齐: 'center'/'left'/'right'/'both'。"""
    ppr = p.find(W + 'pPr')
    if ppr is None:
        ppr = p.makeelement(W + 'pPr', {}); p.insert(0, ppr)
    jc = ppr.find(W + 'jc')
    if jc is None:
        jc = ppr.makeelement(W + 'jc', {}); ppr.append(jc)
    jc.set(qn('w:val'), align)

def _set_run_bold(r):
    """给 run 加粗。"""
    rpr = r.find(W + 'rPr')
    if rpr is None:
        rpr = r.makeelement(W + 'rPr', {}); r.insert(0, rpr)
    if rpr.find(W + 'b') is None:
        rpr.append(rpr.makeelement(W + 'b', {}))

def set_tc_text(tc, text, align=None, bold=False):
    """把单元格(tc)文本设为 text，保留首个 run 的字体格式；支持 \n 换行。
    align: 'center' 等可选对齐；bold: 是否加粗。"""
    # 收集所有段落
    paras = tc.findall(W + 'p')
    if not paras:
        return
    first_p = paras[0]
    # 记录首个 run 的 rPr 作为格式模板
    first_r = first_p.find(W + 'r')
    rpr_tmpl = None
    if first_r is not None:
        rpr = first_r.find(W + 'rPr')
        if rpr is not None:
            rpr_tmpl = clone(rpr)
    # 删除首段之外的段落
    for p in paras[1:]:
        tc.remove(p)
    # 清空首段的 run
    for r in first_p.findall(W + 'r'):
        first_p.remove(r)
    # 记录段落属性 pPr（保留对齐等）
    lines = str(text).split("\n")
    ppr = first_p.find(W + 'pPr')
    target_p = first_p
    for i, line in enumerate(lines):
        if i > 0:
            target_p = clone(first_p)
            if target_p.find(W + 'r') is not None:
                for r in target_p.findall(W + 'r'):
                    target_p.remove(r)
            first_p.addnext(target_p)
        if align:
            _set_para_align(target_p, align)
        r = target_p.makeelement(W + 'r', {})
        if rpr_tmpl is not None:
            r.append(clone(rpr_tmpl))
        force_song5(r)  # 正文/表格填入文字统一宋体五号(封面用set_underline_value,不走这里)
        if bold:
            _set_run_bold(r)
        t = target_p.makeelement(W + 't', {})
        t.set(qn('xml:space'), 'preserve')
        t.text = line
        r.append(t)
        target_p.append(r)

def force_song5(r):
    """把 run 的字体强制为宋体、字号五号(10.5pt=sz21)，保留其它属性(加粗/对齐等)。"""
    rpr = r.find(W + 'rPr')
    if rpr is None:
        rpr = r.makeelement(W + 'rPr', {})
        r.insert(0, rpr)
    fonts = rpr.find(W + 'rFonts')
    if fonts is None:
        fonts = rpr.makeelement(W + 'rFonts', {}); rpr.insert(0, fonts)
    fonts.set(qn('w:hint'), 'eastAsia')
    fonts.set(qn('w:ascii'), '宋体'); fonts.set(qn('w:hAnsi'), '宋体')
    fonts.set(qn('w:eastAsia'), '宋体'); fonts.set(qn('w:cs'), '宋体')
    for tag in ('sz', 'szCs'):
        for e in rpr.findall(W + tag):
            rpr.remove(e)
    sz = rpr.makeelement(W + 'sz', {}); sz.set(qn('w:val'), '21'); rpr.append(sz)
    szcs = rpr.makeelement(W + 'szCs', {}); szcs.set(qn('w:val'), '21'); rpr.append(szcs)

def row_cells(row_el):
    return row_el.findall(W + 'tc')

def tc_text(tc):
    return "".join([(t.text or "") for t in tc.findall('.//' + W + 't')]).strip()
# ---------- 表格行操作 ----------
def tbl_rows(tbl_el):
    return tbl_el.findall(W + 'tr')

def clone_row_after(tbl_el, ref_row):
    """在 ref_row 后克隆一行并返回新行元素（新行内容需自行填充）。"""
    new_row = clone(ref_row)
    ref_row.addnext(new_row)
    return new_row

def clear_cell_images(tc):
    """删除单元格内的所有图片(drawing / VML pict)，保留段落壳。"""
    for tag in (W + 'drawing', ):
        for d in tc.findall('.//' + tag):
            d.getparent().remove(d)
    # VML: w:pict
    for pict in tc.findall('.//' + W + 'pict'):
        pict.getparent().remove(pict)
    # object
    for obj in tc.findall('.//' + W + 'object'):
        obj.getparent().remove(obj)
# ---------- 图片插入（核心：统一 4.8×6.4cm） ----------
from io import BytesIO
from docx.table import Table, _Cell
from PIL import Image as _PILImage, ImageOps as _ImageOps

# 重编码质量（1-95）。高质量、几乎无损；同时清掉手机照片的问题EXIF并修正旋转。
JPEG_QUALITY = 92

def normalize_image(img_path):
    """用 Pillow 规范化图片，返回 (BytesIO, (w,h))。
    - 修正手机 EXIF 旋转方向
    - 清除非标准 EXIF（否则 python-docx 读 DPI 会崩：str __round__）
    - 统一转 RGB JPEG
    """
    im = _PILImage.open(img_path)
    try:
        im = _ImageOps.exif_transpose(im)   # 按 EXIF 摆正方向
    except Exception:
        pass
    if im.mode not in ("RGB", "L"):
        im = im.convert("RGB")
    size = im.size
    bio = BytesIO()
    im.save(bio, format="JPEG", quality=JPEG_QUALITY)
    bio.seek(0)
    return bio, size

def _target_size(size):
    """填满列宽：宽度固定为 IMG_W_CM，高度按原图比例走(不变形)。
    竖图高度超过上限时改按高度反推宽度，避免过高。"""
    try:
        w, h = size
        tw = IMG_W_CM
        th = tw * h / w
        if th > IMG_MAX_H_CM:      # 竖图/过高：按高度封顶反推宽度
            th = IMG_MAX_H_CM
            tw = th * w / h
        return Cm(tw), Cm(th)
    except Exception:
        return Cm(IMG_W_CM), Cm(IMG_H_CM)

def put_picture(tc, doc, img_path, caption=""):
    """在单元格 tc 内放入图片(4.8×6.4)，可选图注。tc 需已在 doc 内。"""
    cell = _Cell(tc, Table(tc.getparent().getparent(), doc))
    # 用单元格首段承载图片，居中
    p = cell.paragraphs[0]
    p.alignment = 1  # center
    run = p.add_run()
    try:
        stream, size = normalize_image(img_path)
        w, h = _target_size(size)
        run.add_picture(stream, width=w, height=h)
    except Exception:
        # 兜底：规范化失败时尝试直接插入原图
        w, h = _target_size(None)
        run.add_picture(img_path, width=w, height=h)
    if caption:
        cp = cell.add_paragraph(caption)
        cp.alignment = 1
        for r in cp.runs:
            force_song5(r._r)

def rebuild_image_table(tbl_el, groups, doc):
    """按 groups=[{title, images:[{path,caption}]}] 重建 2 列图片表。
    使用表内首个整行(合并标题行)与首个图片行作为格式样板。"""
    rows = tbl_rows(tbl_el)
    header_tmpl = clone(rows[0])            # 合并标题行样板（如“试验前图片”）
    img_row_tmpl = clone(rows[1])           # 双图行样板
    # 清空所有现有行
    for r in rows:
        tbl_el.remove(r)

    def new_header(title):
        r = clone(header_tmpl)
        tc = row_cells(r)[0]
        clear_cell_images(tc)
        set_tc_text(tc, title, bold=True)   # 图组标题加粗
        return r

    def _set_row_height(r, cm):
        """给行设最小高度(atLeast)，配合垂直居中形成宽松留白。"""
        trPr = r.find(qn('w:trPr'))
        if trPr is None:
            trPr = r.makeelement(qn('w:trPr'), {})
            r.insert(0, trPr)
        trH = trPr.find(qn('w:trHeight'))
        if trH is None:
            trH = trPr.makeelement(qn('w:trHeight'), {})
            trPr.append(trH)
        trH.set(qn('w:val'), str(int(cm * 567)))   # cm -> twips (1cm≈567)
        trH.set(qn('w:hRule'), 'atLeast')

    def _vcenter(tc):
        """单元格内容垂直居中。"""
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = tc.makeelement(qn('w:tcPr'), {})
            tc.insert(0, tcPr)
        vA = tcPr.find(qn('w:vAlign'))
        if vA is None:
            vA = tcPr.makeelement(qn('w:vAlign'), {})
            tcPr.append(vA)
        vA.set(qn('w:val'), 'center')

    def new_img_row():
        r = clone(img_row_tmpl)
        _set_row_height(r, ROW_H_CM)
        for tc in row_cells(r):
            clear_cell_images(tc)
            set_tc_text(tc, "")   # 清掉供体残留的旧图注
            _vcenter(tc)
        return r

    for g in groups:
        imgs = g.get("images", [])
        if not imgs and not g.get("title"):
            continue
        tbl_el.append(new_header(g.get("title", "")))
        # 每行 2 图
        for i in range(0, len(imgs), 2):
            r = new_img_row()
            cells = row_cells(r)
            for j in range(2):
                if i + j < len(imgs):
                    img = imgs[i + j]
                    put_picture(cells[j], doc, img["path"], img.get("caption", ""))
                else:
                    # 空单元格：清掉图，留空
                    pass
            tbl_el.append(r)
# ---------- 测试段落（从供体克隆一整段并填充） ----------
def _donor_section_blocks(donor_doc):
    """从供体取“高温耐久测试”整段的块序列（标题→…→试验图片表）作为样板。
    返回块元素列表（deepcopy 后可独立使用）。"""
    body = donor_doc.element.body
    els = list(body.iterchildren())
    start = end = None
    for i, el in enumerate(els):
        if el.tag == W + 'p' and para_text(el) == "高温耐久测试":
            start = i
        if el.tag == W + 'p' and para_text(el) == "振动测试":
            end = i
            break
    blocks = [clone(e) for e in els[start:end]]
    return blocks

def _section_tables(blocks):
    """返回该段块列表中的表格元素，按出现顺序：
    [样品信息, 试验设备, 试验描述, 试验结论, 试验图片]"""
    tbls = [b for b in blocks if b.tag == W + 'tbl']
    return tbls

def fill_sample_tbl(tbl, doc, d):
    # 1行4列：样品名称|<值>|样品编号|<值>
    tc = row_cells(tbl_rows(tbl)[0])
    set_tc_text(tc[1], d.get("sample_name", ""))
    set_tc_text(tc[3], d.get("sample_no", ""))

def fill_equip_tbl(tbl, doc, equlist):
    # 表头行R0固定；数据行按 equlist 生成（序号,名称,型号,管理编号,校准有效期）
    rows = tbl_rows(tbl)
    header, data_tmpl = rows[0], rows[1]
    # 删除除表头外所有行
    for r in rows[1:]:
        tbl.remove(r)
    for i, e in enumerate(equlist, 1):
        r = clone(data_tmpl)
        c = row_cells(r)
        vals = [str(i), e.get("name", ""), e.get("model", ""), e.get("mgmt_no", ""), e.get("cal_valid", "")]
        for k, tc in enumerate(c[:5]):
            clear_cell_images(tc)
            set_tc_text(tc, vals[k])
        tbl.append(r)

def fill_desc_tbl(tbl, doc, d):
    # 4行：环境条件|<>|测试日期|<> ; 试验标准|<> ; 试验条件|<>(+可选配图) ; 试验要求|<>
    rows = tbl_rows(tbl)
    r0 = row_cells(rows[0]); set_tc_text(r0[1], d.get("env", ""), align="center"); set_tc_text(r0[3], d.get("test_date", ""))
    r1 = row_cells(rows[1]); set_tc_text(r1[1], d.get("standard", ""))
    r2 = row_cells(rows[2]); clear_cell_images(r2[1]); set_tc_text(r2[1], d.get("condition", ""))
    r3 = row_cells(rows[3]); set_tc_text(r3[1], d.get("requirement", ""))
    # 试验条件配图：插到条件文字下方（延后到表已入 doc 时执行，见 build_section）

def fill_concl_tbl(tbl, doc, samples):
    # R0表头;每个样品一行(编号|结果|结论);末行备注
    rows = tbl_rows(tbl)
    header, data_tmpl, remark = rows[0], rows[1], rows[-1]
    for r in rows[1:-1]:
        tbl.remove(r)
    for s in samples:
        r = clone(data_tmpl)
        c = row_cells(r)
        set_tc_text(c[0], s.get("no", ""))
        set_tc_text(c[1], s.get("result", ""))
        set_tc_text(c[2], s.get("conclusion", ""))
        remark.addprevious(r)

def build_section(doc, donor_doc, test):
    """在 doc 的“报告结束”段前插入一个填充好的测试段落。test 为该测试项的数据字典。"""
    blocks = _donor_section_blocks(donor_doc)
    tbls = _section_tables(blocks)
    # 标题（blocks[0]）
    set_tc_text  # noqa
    title_p = blocks[0]
    # 用第一处 run 改标题文本
    set_para_text(title_p, test.get("title", "测试"))
    # 每个测试项目单独分页：标题段前加分页
    add_page_break_before(title_p)
    # 填表
    fill_sample_tbl(tbls[0], doc, test)
    fill_equip_tbl(tbls[1], doc, test.get("equipment", []))
    fill_desc_tbl(tbls[2], doc, test)
    fill_concl_tbl(tbls[3], doc, test.get("samples", []))
    # 图片表先插入文档再重建（图片关系需绑定到 doc）
    end_p = find_para(doc, "报告结束")
    anchor = end_p
    inserted = []
    for b in blocks:
        anchor.addprevious(b)
        inserted.append(b)
    # 试验条件配图：表已入 doc，可绑定图片关系。插到"试验条件"单元格文字下方
    cond_imgs = test.get("condition_images", [])
    if cond_imgs:
        desc_rows = tbl_rows(tbls[2])
        cond_cell = row_cells(desc_rows[2])[1]
        append_pictures_to_cell(cond_cell, doc, cond_imgs)
    # 重建图片表（此时已在 doc 内）
    img_tbl = tbls[4]
    rebuild_image_table(img_tbl, test.get("image_groups", []), doc)
    return inserted

def append_pictures_to_cell(tc, doc, imgs):
    """在单元格已有文字下方，逐张追加居中图片(带可选图注)，不覆盖原文字。"""
    cell = _Cell(tc, Table(tc.getparent().getparent(), doc))
    for im in imgs:
        p = cell.add_paragraph()
        p.alignment = 1  # center
        run = p.add_run()
        try:
            stream, size = normalize_image(im["path"])
            w, h = _target_size(size)
            run.add_picture(stream, width=w, height=h)
        except Exception:
            w, h = _target_size(None)
            run.add_picture(im["path"], width=w, height=h)
        cap = im.get("caption", "")
        if cap:
            cp = cell.add_paragraph(cap)
            cp.alignment = 1
            for r in cp.runs:
                force_song5(r._r)

def add_page_break_before(p):
    """给段落加 pageBreakBefore，使其从新页开始。"""
    ppr = p.find(W + 'pPr')
    if ppr is None:
        ppr = p.makeelement(W + 'pPr', {})
        p.insert(0, ppr)
    if ppr.find(W + 'pageBreakBefore') is None:
        pbb = p.makeelement(W + 'pageBreakBefore', {})
        ppr.insert(0, pbb)

def _run_has_media(r):
    """run 内是否含图片/图形/对象(不能删)。"""
    return (r.find('.//' + W + 'drawing') is not None
            or r.find('.//' + W + 'pict') is not None
            or r.find('.//' + W + 'object') is not None)

def set_para_text(p, text):
    """设置段落文本，保留首个文本run格式；保留含图片/图形的run(不误删页眉logo等)。"""
    runs = p.findall(W + 'r')
    text_runs = [r for r in runs if not _run_has_media(r)]
    rpr_tmpl = None
    if text_runs:
        rpr = text_runs[0].find(W + 'rPr')
        if rpr is not None:
            rpr_tmpl = clone(rpr)
    # 只删文本run，保留图片run
    first_text_run = text_runs[0] if text_runs else None
    for r in text_runs:
        p.remove(r)
    r = p.makeelement(W + 'r', {})
    if rpr_tmpl is not None:
        r.append(clone(rpr_tmpl))
    t = p.makeelement(W + 't', {})
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    # 尽量放回原文本位置：插到第一个图片run之后或段末
    media_runs = [x for x in p.findall(W + 'r') if _run_has_media(x)]
    if media_runs:
        media_runs[-1].addnext(r)
    else:
        p.append(r)
# ---------- 封面 / 样品信息 / 汇总 / 页眉 ----------
def set_underline_value(tc, text):
    """填写"填空式"单元格：保留底部横线(由带下划线的空格run构成)，
    只把值填进首个文字run并加下划线，横线因此延续到值文字下方。"""
    paras = tc.findall(W + 'p')
    if not paras:
        set_tc_text(tc, text); return
    p = paras[0]
    runs = p.findall(W + 'r')
    if not runs:
        set_tc_text(tc, text); return
    # 找带下划线的空格run作为"横线"样板；文字run放值
    value_run = runs[0]
    # 收集其余run里"带下划线的空白run"（横线），删掉其它含文字的run
    keep_fillers = []
    for r in runs[1:]:
        t = "".join([(x.text or "") for x in r.findall(W + 't')])
        has_u = r.find('.//' + W + 'u') is not None
        if t.strip() == "" and has_u:
            keep_fillers.append(r)  # 横线空格，保留
        else:
            p.remove(r)             # 旧占位文字，删除
    # 设置值文字：强制值文字本身不带下划线（显式 u=none 以覆盖样式继承来的下划线）
    # 只保留后面空格run组成的底部大横线。
    for x in value_run.findall(W + 't'):
        value_run.remove(x)
    rpr = value_run.find(W + 'rPr')
    if rpr is None:
        rpr = value_run.makeelement(W + 'rPr', {}); value_run.insert(0, rpr)
    for u in rpr.findall(W + 'u'):
        rpr.remove(u)
    u = rpr.makeelement(W + 'u', {}); u.set(qn('w:val'), 'none'); rpr.append(u)
    t = value_run.makeelement(W + 't', {}); t.set(qn('xml:space'), 'preserve'); t.text = text
    value_run.append(t)

def fill_cover(doc, info):
    """封面表(TBL#0)：R0样品名称 R1零件号 R2型号 R3验证阶段（值在第1列）。"""
    tbl = doc.tables[0]
    rows = tbl_rows(tbl._tbl)
    mapping = {0: "sample_name", 1: "sample_no", 2: "sample_model", 3: "verify_phase"}
    for ri, key in mapping.items():
        cells = row_cells(rows[ri])
        if len(cells) > 1 and info.get(key):
            set_underline_value(cells[1], info.get(key, ""))

def fill_sample_info(doc, info):
    """样品信息表(TBL#1)，13行。按标签匹配填第2列/第4列。"""
    tbl = doc.tables[1]
    label_key = {
        "委托方名称": "client_name", "委托方地址": "client_addr",
        "制造商名称": "maker_name", "制造商地址": "maker_addr",
        "样品名称": "sample_name", "样品零件号": "sample_no",
        "样品型号": "sample_model", "样品数量": "sample_qty",
        "来样方式": "sample_way", "收样日期": "recv_date",
        "委托单号": "commission_no", "检测日期": "test_date_range",
        "检测单位": "lab_name", "检测项目": "test_items",
        "检测依据": "test_basis", "备注": "remark",
    }
    for row in tbl_rows(tbl._tbl):
        cells = row_cells(row)
        # 第1列标签->第2列值
        if len(cells) >= 2:
            lab = tc_text(cells[0])
            key = label_key.get(lab)
            if key and info.get(key) is not None and info.get(key) != "":
                set_tc_text(cells[1], info.get(key))
        # 第3列标签->第4列值
        if len(cells) >= 4:
            lab2 = tc_text(cells[2])
            key2 = label_key.get(lab2)
            if key2 and info.get(key2) is not None and info.get(key2) != "":
                set_tc_text(cells[3], info.get(key2))

def derive_result(t):
    """汇总"试验结果"：优先用手填 overall_result；否则按各样品结论自动判定。"""
    v = (t.get("overall_result") or "").strip()
    if v:
        return v
    concls = [(s.get("conclusion") or "").strip() for s in t.get("samples", [])]
    concls = [c for c in concls if c]
    if not concls:
        return "合格"
    return "合格" if all(c == "合格" for c in concls) else "不合格"

def fill_summary(doc, tests):
    """结果汇总表(TBL#2)：R0/R1为表头，数据行按测试项生成。
    列：序号|试验项目|标准号/条款号|样机编号|开始时间|完成时间|试验结果。"""
    tbl = doc.tables[2]
    rows = tbl_rows(tbl._tbl)
    header_rows = rows[:2]
    data_tmpl = rows[2]
    for r in rows[2:]:
        tbl._tbl.remove(r)
    for i, t in enumerate(tests, 1):
        r = clone(data_tmpl)
        c = row_cells(r)
        vals = [str(i), t.get("title", ""), t.get("standard", ""),
                t.get("sample_no", ""), t.get("start_date", ""),
                t.get("end_date", ""), derive_result(t)]
        for k, tc in enumerate(c[:7]):
            clear_cell_images(tc)
            set_tc_text(tc, vals[k])
        tbl._tbl.append(r)

def fill_report_no(doc, report_no):
    """更新正文报告编号段 + 页眉报告编号。"""
    if not report_no:
        return
    for p in doc.element.body.iterchildren():
        if p.tag == W + 'p' and para_text(p).startswith("报告编号"):
            set_para_text(p, "报告编号：" + report_no)
            break
    # 页眉：报告编号靠右上角（logo保留在左，编号用右tab推到右边缘）
    for section in doc.sections:
        for hp in section.header.paragraphs:
            if "报告编号" in hp.text:
                set_header_reportno_right(hp._p, "报告编号：" + report_no)

def set_header_reportno_right(p, text):
    """页眉段：报告编号靠最右侧，宋体小四加粗。logo为浮动图片(不占行宽)，
    故直接把段落设为右对齐即可，保留图片run。"""
    # 1) 删除所有纯文本run（保留含图片/图形的run）
    for r in p.findall(W + 'r'):
        if not _run_has_media(r):
            p.remove(r)
    # 2) 段落右对齐，清掉旧tabs（避免落在页眉样式的居中tab上）
    ppr = p.find(W + 'pPr')
    if ppr is None:
        ppr = p.makeelement(W + 'pPr', {})
        p.insert(0, ppr)
    for old in ppr.findall(W + 'tabs'):
        ppr.remove(old)
    jc = ppr.find(W + 'jc')
    if jc is None:
        jc = ppr.makeelement(W + 'jc', {}); ppr.append(jc)
    jc.set(qn('w:val'), 'right')
    # 3) 新建文字run：宋体 + 加粗 + 小四(sz24)
    r = p.makeelement(W + 'r', {})
    rpr = r.makeelement(W + 'rPr', {})
    fonts = rpr.makeelement(W + 'rFonts', {})
    fonts.set(qn('w:hint'), 'eastAsia')
    fonts.set(qn('w:ascii'), '宋体'); fonts.set(qn('w:hAnsi'), '宋体'); fonts.set(qn('w:eastAsia'), '宋体')
    rpr.append(fonts)
    b = rpr.makeelement(W + 'b', {}); rpr.append(b)
    sz = rpr.makeelement(W + 'sz', {}); sz.set(qn('w:val'), '24'); rpr.append(sz)
    szcs = rpr.makeelement(W + 'szCs', {}); szcs.set(qn('w:val'), '24'); rpr.append(szcs)
    r.append(rpr)
    t = r.makeelement(W + 't', {}); t.set(qn('xml:space'), 'preserve'); t.text = text
    r.append(t)
    # 放到最后一个图片run之后
    media_runs = [x for x in p.findall(W + 'r') if _run_has_media(x)]
    if media_runs:
        media_runs[-1].addnext(r)
    else:
        p.append(r)

def build_dynamic_toc(doc):
    """把写死的目录项替换为 Word 目录域（TOC field），自动抓取所有一级标题。
    页码在生成后由 COM 刷新域填入。"""
    body = doc.element.body
    els = list(body.iterchildren())
    # 找"目录"标题段
    toc_head = None
    for el in els:
        if el.tag == W + 'p' and para_text(el) == "目录":
            toc_head = el
            break
    if toc_head is None:
        return
    # 删除其后连续的 pStyle=7 目录项段
    nxt = toc_head.getnext()
    removed_any = False
    while nxt is not None and nxt.tag == W + 'p':
        ppr = nxt.find(W + 'pPr')
        ps = ppr.find(W + 'pStyle') if ppr is not None else None
        if ps is not None and ps.get(qn('w:val')) == "7":
            to_del = nxt
            nxt = nxt.getnext()
            body.remove(to_del)
            removed_any = True
        else:
            break
    # 构造 TOC 域段落，插入到"目录"标题之后
    p = toc_head.makeelement(W + 'p', {})
    ppr = p.makeelement(W + 'pPr', {})
    ps = p.makeelement(W + 'pStyle', {}); ps.set(qn('w:val'), "7")
    ppr.append(ps)
    tabs = p.makeelement(W + 'tabs', {})
    tab = p.makeelement(W + 'tab', {}); tab.set(qn('w:val'), "right"); tab.set(qn('w:leader'), "dot"); tab.set(qn('w:pos'), "9070")
    tabs.append(tab); ppr.append(tabs)
    p.append(ppr)

    def run_with(child):
        r = p.makeelement(W + 'r', {})
        r.append(child)
        return r
    fb = p.makeelement(W + 'fldChar', {}); fb.set(qn('w:fldCharType'), "begin")
    fb.set(qn('w:dirty'), "true")   # 标记目录域为“脏”，WPS/Word 打开时会自动重算页码
    p.append(run_with(fb))
    it = p.makeelement(W + 'instrText', {}); it.set(qn('xml:space'), "preserve"); it.text = ' TOC \\o "1-1" \\h \\z \\u '
    p.append(run_with(it))
    fs = p.makeelement(W + 'fldChar', {}); fs.set(qn('w:fldCharType'), "separate")
    p.append(run_with(fs))
    t = p.makeelement(W + 't', {}); t.set(qn('xml:space'), "preserve"); t.text = "（打开文档时选“是/更新”即自动生成页码；若未提示，右键此处选“更新域”）"
    p.append(run_with(t))
    fe = p.makeelement(W + 'fldChar', {}); fe.set(qn('w:fldCharType'), "end")
    p.append(run_with(fe))
    toc_head.addnext(p)

def set_update_fields_on_open(doc):
    """在文档设置里写入 <w:updateFields val="true"/>，
    让 Word/WPS 打开文档时自动更新所有域（含目录页码）。
    这是跨平台方案：不依赖服务器装 Office/COM，用户一打开即自动刷新。"""
    settings = doc.settings.element
    # 已存在就复用，避免重复
    uf = settings.find(W + 'updateFields')
    if uf is None:
        uf = settings.makeelement(W + 'updateFields', {})
        # 按 OOXML schema，updateFields 应排在 settings 靠前位置，插到最前最稳妥
        settings.insert(0, uf)
    uf.set(qn('w:val'), "true")
# ---------- 主入口 ----------
def generate(project, out_path):
    """project: {info:{...}, tests:[{...}, ...]}  ->  生成 out_path (.docx)"""
    info = project.get("info", {})
    tests = project.get("tests", [])
    doc = Document(SKELETON)
    donor = Document(DONOR)

    fill_report_no(doc, info.get("report_no", ""))
    fill_cover(doc, info)
    fill_sample_info(doc, info)
    fill_summary(doc, tests)
    build_dynamic_toc(doc)

    for t in tests:
        # 段内"样品名称"始终与首页(样品信息)一致
        t["sample_name"] = info.get("sample_name", "")
        build_section(doc, donor, t)

    # 让 Word/WPS 打开文档时自动更新目录页码（跨平台，不依赖服务器 Office）
    set_update_fields_on_open(doc)

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    # 生成后把目录页码算好写死进文件，用户打开即成品：
    #  - Windows(开发机)：用 WPS/Word COM
    #  - Linux(服务器)：用 LibreOffice 无头模式
    # 两者都失败也不影响出报告，文件里已带 updateFields/dirty 标记作兜底。
    refresh_toc_and_fields(out_path)
    return out_path

def refresh_toc_and_fields(path):
    """把目录页码算好写死进 docx。全部失败也不抛异常（报告照常产出，
    文件已带 updateFields/dirty 标记，用户可手动更新兜底）。"""
    import sys as _sys
    # 1) 本机就是 Windows(开发机/单机部署)：直接 WPS/Word COM
    if _sys.platform.startswith("win"):
        try:
            update_fields_com(path)
        except Exception:
            pass
        return
    # 2) Linux 服务器：把 docx 发给 Windows 上的“目录更新服务”刷新（秒级）
    #    配置见项目根目录 toc_service.json（没配置就跳过，不再走 4 分钟的 LibreOffice）
    try:
        update_toc_via_service(path)
    except Exception:
        pass

# 默认指向那台装了 WPS 的 Windows（192.168.24.68）。
# 换机器 IP 时，改这里，或在项目根放 toc_service.json / 设环境变量覆盖，都行。
DEFAULT_TOC_SERVICE_URL = "http://192.168.24.68:8765"

def _load_toc_service_cfg():
    """确定目录更新服务地址。优先级：环境变量 > toc_service.json > 内置默认。
    默认已内置，所以 VM 上零配置即可用；连不上会静默跳过。"""
    import json
    cfg = {}
    fp = os.path.join(BASE, "toc_service.json")
    if os.path.exists(fp):
        try:
            with open(fp, "r", encoding="utf-8") as f:
                cfg = json.load(f) or {}
        except Exception:
            cfg = {}
    if os.environ.get("TOC_SERVICE_URL"):
        cfg["url"] = os.environ["TOC_SERVICE_URL"]
    if os.environ.get("TOC_SERVICE_TOKEN"):
        cfg["token"] = os.environ["TOC_SERVICE_TOKEN"]
    url = (cfg.get("url") or DEFAULT_TOC_SERVICE_URL or "").strip().rstrip("/")
    if not url:
        return None
    return {"url": url, "token": cfg.get("token", ""),
            "timeout": int(cfg.get("timeout", 90))}

def update_toc_via_service(path):
    """把 docx POST 给 Windows 目录更新服务，用返回的成品覆盖原文件。
    用 urllib(标准库)，不依赖 requests。连不上/失败都静默跳过。"""
    cfg = _load_toc_service_cfg()
    if not cfg:
        return  # 没配置服务，跳过
    import urllib.request
    with open(path, "rb") as f:
        body = f.read()
    req = urllib.request.Request(cfg["url"] + "/update_toc", data=body, method="POST")
    req.add_header("Content-Type", "application/octet-stream")
    if cfg["token"]:
        req.add_header("X-Token", cfg["token"])
    with urllib.request.urlopen(req, timeout=cfg["timeout"]) as resp:
        out = resp.read()
    # 只有拿到看起来正常的 docx(zip 头 PK)才覆盖，避免把错误页写进文件
    if out[:2] == b"PK" and len(out) > 1000:
        with open(path, "wb") as f:
            f.write(out)

def update_toc_libreoffice(path):
    """用 LibreOffice 无头模式刷新目录/域。需要系统装 libreoffice + python 能 import uno。"""
    import subprocess, shutil
    script = os.path.join(os.path.dirname(os.path.abspath(__file__)), "toc_lo.py")
    if not os.path.exists(script):
        return
    soffice = None
    for name in ("soffice", "libreoffice"):
        soffice = shutil.which(name)
        if soffice:
            break
    if not soffice:
        for p in ("/usr/bin/soffice", "/usr/bin/libreoffice",
                  "/opt/libreoffice/program/soffice", "/snap/bin/libreoffice"):
            if os.path.exists(p):
                soffice = p; break
    if not soffice:
        return  # 没装 LibreOffice，静默跳过（靠 updateFields/dirty 兜底）
    # 找一个能 import uno 的 python：优先系统 python3（装了 python3-uno），
    # 再试 LibreOffice 自带 python，最后退回当前解释器。
    candidates = ["/usr/bin/python3", "python3",
                  "/usr/lib/libreoffice/program/python",
                  "/opt/libreoffice/program/python"]
    import sys as _sys
    candidates.append(_sys.executable)
    py = None
    for c in candidates:
        # c 含路径分隔符 → 当成绝对/相对路径，必须存在；否则当成命令名去 PATH 里找
        if os.path.sep in c or "/" in c:
            exe = c if os.path.exists(c) else None
        else:
            exe = shutil.which(c)
        if not exe:
            continue
        try:
            r = subprocess.run([exe, "-c", "import uno"], capture_output=True, timeout=20)
            if r.returncode == 0:
                py = exe; break
        except Exception:
            continue
    if not py:
        return  # 没有能用的 uno，跳过
    try:
        subprocess.run([py, script, os.path.abspath(path), soffice],
                       capture_output=True, timeout=180)
    except Exception:
        pass

def update_fields_com(path):
    """用 WPS/Word COM 打开文档，更新所有域(含目录)与页码后保存。
    在后台线程运行，需要 CoInitialize。"""
    import pythoncom
    path = os.path.abspath(path)
    pythoncom.CoInitialize()
    app = None
    try:
        import win32com.client as win32
        try:
            app = win32.gencache.EnsureDispatch("KWps.Application")   # 优先 WPS
        except Exception:
            app = win32.gencache.EnsureDispatch("Word.Application")   # 退回 Word
        try:
            app.Visible = False
        except Exception:
            pass
        try:
            app.DisplayAlerts = False
        except Exception:
            pass
        doc = app.Documents.Open(path)
        # 更新目录（TOC）
        try:
            for toc in doc.TablesOfContents:
                toc.Update()
        except Exception:
            pass
        # 更新正文所有域
        try:
            doc.Fields.Update()
        except Exception:
            pass
        # 更新页眉页脚域
        try:
            for sec in doc.Sections:
                for hf in (sec.Headers, sec.Footers):
                    for h in hf:
                        h.Range.Fields.Update()
        except Exception:
            pass
        doc.Save()
        doc.Close(False)
    finally:
        try:
            if app is not None:
                app.Quit()
        except Exception:
            pass
        pythoncom.CoUninitialize()

