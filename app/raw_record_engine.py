# -*- coding: utf-8 -*-
"""原始记录生成引擎：从「原始记录骨架.docx」按测试项逐个生成原始记录。

每个试验项目生成一张原始记录表（15 行 4 列）：
- 样品基础信息来自首页 info（可由申请单自动导入）
- 仪器设备复用设备库（test.equipment）
- 试验项目/标准/条件/要求复用试验项目库（test）

多个测试项时，每张记录之间插入分页符，一份文档包含全部原始记录。
"""
import os
import copy
from docx import Document
from docx.oxml.ns import qn

# 复用主报告引擎的低层助手，保证字体/换行处理一致（宋体五号、\r\n 归一等）
import report_engine as E

W = E.W

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SKELETON = os.path.join(BASE, "模板库", "原始记录骨架.docx")


def _row_cells(row_el):
    return row_el.findall(W + 'tc')


def _set(tc, text):
    """填单元格文本，沿用主引擎的换行/字体处理。"""
    E.set_tc_text(tc, "" if text is None else str(text))


def _join(items):
    """多台设备等多值信息按换行拼接（模板里就是多行显示）。"""
    return "\n".join(s for s in items if str(s or "").strip())


def _fill_table(tbl_el, info, test, doc):
    """填一张原始记录表（15 行）。tbl_el 为 <w:tbl> 元素。"""
    # 移除表格浮动属性，让表格可以跨页
    tblPr = tbl_el.find(qn('w:tblPr'))
    if tblPr is not None:
        tblpPr = tblPr.find(qn('w:tblpPr'))
        if tblpPr is not None:
            tblPr.remove(tblpPr)

    rows = tbl_el.findall(W + 'tr')

    def cells(ri):
        return _row_cells(rows[ri])

    equ = test.get("equipment", []) or []
    eq_names = _join([e.get("name", "") for e in equ])
    eq_models = _join([e.get("model", "") for e in equ])
    eq_nos = _join([e.get("mgmt_no", "") for e in equ])
    eq_cal = _join([e.get("cal_valid", "") for e in equ])

    # R0 样品名称 / 样品型号
    c = cells(0); _set(c[1], info.get("sample_name", "")); _set(c[3], info.get("sample_model", ""))
    # R1 样品零件号 / 委托单号
    c = cells(1); _set(c[1], info.get("sample_no", "")); _set(c[3], info.get("commission_no", ""))
    # R2 样品数量 / 样品编号（样品编号取该测试项的样机编号）
    c = cells(2); _set(c[1], info.get("sample_qty", "")); _set(c[3], test.get("sample_no", ""))
    # R3 额定电压 / 环境条件（额定电压有值且未带单位时自动补 V）
    c = cells(3); _set(c[1], _with_volt_unit(info.get("rated_volt", ""))); _set(c[3], test.get("env", ""))
    # R4 设备名称 / 设备编号
    c = cells(4); _set(c[1], eq_names); _set(c[3], eq_nos)
    # R5 设备型号 / 校准周期
    c = cells(5); _set(c[1], eq_models); _set(c[3], eq_cal)
    # R6 试验项目 / 试验标准
    c = cells(6); _set(c[1], test.get("title", "")); _set(c[3], test.get("standard", ""))
    # R7 试验条件（跨列）：文字上下居中，其下追加试验条件配图
    _cond_text = test.get("condition", "")
    _cond_imgs = test.get("condition_images", [])
    c = cells(7); _set(c[1], _cond_text); _vcenter(c[1])
    _append_cond_images(c[1], doc, _cond_imgs)
    # 仅当试验条件内容较多（有配图或文字超阈值）时，才分页+撕框；
    # 内容少时保持默认行高、不插分页符，避免产生空白页。
    _need_split = bool(_cond_imgs) or _condition_is_long(_cond_text, _DEFAULT_ROW_HEIGHTS.get(7))
    # 分页模式下撑高 R7，填满第一页
    if _need_split:
        h7 = _page1_r7_height(doc, rows)
        if h7 > 0:
            _row_min_height(rows[7], h7)
    c = cells(8); _set(c[1], test.get("requirement", ""))
    if _need_split:
        _page_break_before_cell(c[0])
    # R10 试验状态：仅分页模式下才拉高
    # 不再撑高 R10，让表格自然排版
    if _need_split:
        h10 = _page2_status_height(doc, rows)
        if h10 > 0:
            _row_min_height(rows[10], h10)
    # R13 测试日期：内容居中放（标签+值单元格都水平+垂直居中）；值仍留空白由用户手填
    c = cells(13); _center_cell(c[0]); _center_cell(c[1])


def _vcenter(tc):
    """单元格内容垂直居中（试验条件文字上下居中）。"""
    from docx.oxml.ns import qn
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = tc.makeelement(qn('w:tcPr'), {})
        tc.insert(0, tcPr)
    vA = tcPr.find(qn('w:vAlign'))
    if vA is None:
        vA = tcPr.makeelement(qn('w:vAlign'), {})
        tcPr.append(vA)
    vA.set(qn('w:val'), 'center')


def _center_cell(tc):
    """单元格内容水平+垂直居中（每个段落设 jc=center，单元格设 vAlign=center）。"""
    from docx.oxml.ns import qn
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = tc.makeelement(qn('w:tcPr'), {})
        tc.insert(0, tcPr)
    vA = tcPr.find(qn('w:vAlign'))
    if vA is None:
        vA = tcPr.makeelement(qn('w:vAlign'), {})
        tcPr.append(vA)
    vA.set(qn('w:val'), 'center')
    for p in tc.findall(qn('w:p')):
        ppr = p.find(qn('w:pPr'))
        if ppr is None:
            ppr = p.makeelement(qn('w:pPr'), {})
            p.insert(0, ppr)
        jc = ppr.find(qn('w:jc'))
        if jc is None:
            jc = ppr.makeelement(qn('w:jc'), {})
            ppr.append(jc)
        jc.set(qn('w:val'), 'center')


def _row_min_height(row_el, twips):
    """给行设最小行高（atLeast），使单元格纵向拉高填满页面。twips：1cm≈567。"""
    trPr = row_el.find(qn('w:trPr'))
    if trPr is None:
        trPr = row_el.makeelement(qn('w:trPr'), {})
        row_el.insert(0, trPr)
    trH = trPr.find(qn('w:trHeight'))
    if trH is None:
        trH = trPr.makeelement(qn('w:trHeight'), {})
        trPr.append(trH)
    trH.set(qn('w:val'), str(int(twips)))
    trH.set(qn('w:hRule'), 'atLeast')


def _usable_text_height(doc):
    """正文可用高度(twips) = 页高 - 上下页边距。"""
    sec = doc.sections[0]
    try:
        return int(sec.page_height.twips) - int(sec.top_margin.twips) - int(sec.bottom_margin.twips)
    except Exception:
        return 14500  # A4 常见值兜底


# 模板各行默认行高(twips)，当行元素没有显式 w:trHeight 时用作兆底
_DEFAULT_ROW_HEIGHTS = {
    0: 510, 1: 510, 2: 510, 3: 510,
    4: 637, 5: 680, 6: 510, 7: 2283,
    8: 631, 9: 567, 10: 3331, 11: 601,
    12: 400, 13: 417, 14: 455,
}


def _rows_height_sum(rows, idxs):
    """若干行的最小行高(trHeight)之和(twips)。
    有显式 w:trHeight 用显式值；否则回退到模板默认行高。"""
    tot = 0
    for i in idxs:
        val = None
        trPr = rows[i].find(qn('w:trPr'))
        if trPr is not None:
            trH = trPr.find(qn('w:trHeight'))
            if trH is not None and trH.get(qn('w:val')):
                try:
                    val = int(trH.get(qn('w:val')))
                except Exception:
                    pass
        if val is None:
            val = _DEFAULT_ROW_HEIGHTS.get(i, 0)
        tot += val
    return tot


# 标题(试验记录表, sz30)+编号行(sz24)+段间距，实测约占这么多 twips；
# 加一点安全余量，宁可 R7 稍矮留白，也不要溢出把整行挤到下一页。
_HEADER_PARA_TWIPS = 1400
_PAGE1_SAFETY_TWIPS = 1200


# R7 试验条件值单元格跨3列(宽8391twips≈14.8cm)，9pt宋体一行约放46个中文字。
# 按「字宽」计(中文/全角=2, ASCII=1)，一行约 92 个字宽，留点余量取 88。
_COND_CHARS_PER_LINE = 88
_COND_LINE_TWIPS = 260      # 9pt 单行约占高度(twips)
_COND_CELL_PAD_TWIPS = 300  # 单元格上下内边距余量


def _visual_line_count(text, chars_per_line=_COND_CHARS_PER_LINE):
    """估算文字实际渲染的可视行数：每个显式换行至少占 1 行，
    过长的行按每行 chars_per_line 个「字宽」自动折行(中文/全角=2, ASCII=1)。"""
    if not text:
        return 0
    total = 0
    lines = str(text).replace('\r\n', '\n').replace('\r', '\n').split('\n')
    for line in lines:
        width = sum(2 if ord(ch) >= 0x2E80 else 1 for ch in line)
        total += max(1, -(-width // chars_per_line))  # ceil 除法
    return total


def _condition_is_long(text, r7_default=None):
    """判断试验条件是否多到放不进默认行高、需要分页。
    按可视行数估算所需高度，只有超过 R7 默认行高时才判定为「长」，
    避免几行短句(换行多但内容少)被误判成需要分页。"""
    if not text:
        return False
    if r7_default is None:
        r7_default = _DEFAULT_ROW_HEIGHTS.get(7, 2283)
    needed = _visual_line_count(text) * _COND_LINE_TWIPS + _COND_CELL_PAD_TWIPS
    return needed > r7_default






def _page1_r7_height(doc, rows):
    """R7(试验条件)在第一页能占的最大行高：
    可用高度 - 标题/编号 - 信息行(R0~R6) - 安全余量。
    这样 R7 拉到最大又不会脱离第一页。"""
    usable = _usable_text_height(doc)
    info_h = _rows_height_sum(rows, range(0, 7))
    h = usable - _HEADER_PARA_TWIPS - info_h - _PAGE1_SAFETY_TWIPS
    return max(0, h)


def _page2_status_height(doc, rows):
    """R10(试验状态)在分页模式下的合理高度。
    不撑满整页（会把 R11~R14 挤走产生空白），而是固定一个较大值。"""
    # 固定拉高到约 6000 twips (约 10.5cm)，足够容纳较多内容，
    # 又不至于把签字行挤到单独页造成空白。
    return 6000


def _page_break_before_cell(tc):
    """给单元格首段加 pageBreakBefore，使该行从新一页开始。"""
    p = tc.find(qn('w:p'))
    if p is None:
        return
    ppr = p.find(qn('w:pPr'))
    if ppr is None:
        ppr = p.makeelement(qn('w:pPr'), {})
        p.insert(0, ppr)
    if ppr.find(qn('w:pageBreakBefore')) is None:
        ppr.insert(0, ppr.makeelement(qn('w:pageBreakBefore'), {}))


def _append_cond_images(tc, doc, imgs):
    """在试验条件文字下方追加配图：每行横排 2 张，避免竖排把版面撑乱。
    在单元格内嵌一张 2 列表格承载图片（无边框、居中）。"""
    if not imgs:
        return
    from docx.table import Table, _Cell
    cell = _Cell(tc, Table(tc.getparent().getparent(), doc))
    nrows = (len(imgs) + 1) // 2
    nested = cell.add_table(rows=nrows, cols=2)
    try:
        nested.autofit = True
    except Exception:
        pass
    _no_borders(nested._tbl)
    for idx, im in enumerate(imgs):
        r, c = idx // 2, idx % 2
        ncell = nested.cell(r, c)
        p = ncell.paragraphs[0]
        p.alignment = 1  # center
        run = p.add_run()
        try:
            stream, size = E.normalize_image(im["path"])
            w, h = E._target_size(size)
            run.add_picture(stream, width=w, height=h)
        except Exception:
            w, h = E._target_size(None)
            run.add_picture(im["path"], width=w, height=h)
        cap = im.get("caption", "")
        if cap:
            cp = ncell.add_paragraph(cap)
            cp.alignment = 1
            for rr in cp.runs:
                E.force_song5(rr._r)


def _no_borders(tbl_el):
    """把内嵌表格的四周及内部边框全部设为 none。"""
    tblPr = tbl_el.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = tbl_el.makeelement(qn('w:tblPr'), {})
        tbl_el.insert(0, tblPr)
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is None:
        borders = tblPr.makeelement(qn('w:tblBorders'), {})
        tblPr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = borders.find(qn('w:' + edge))
        if e is None:
            e = borders.makeelement(qn('w:' + edge), {})
            borders.append(e)
        e.set(qn('w:val'), 'none')
        e.set(qn('w:sz'), '0')
        e.set(qn('w:space'), '0')


def _with_volt_unit(v):
    """额定电压：有值且结尾未带 V/v 时补上单位 V。"""
    s = str(v or "").strip()
    if not s:
        return s
    if s[-1] in ("V", "v"):
        return s
    return s + "V"


def _raw_record_no(no):
    """原始记录编号：把委托单号的 ME/WTD（含全角 ME／WTD）前缀替换为 SY。"""
    s = str(no or "").strip()
    for pref in ("ME/WTD", "ME／WTD"):
        if s.upper().startswith(pref.upper()):
            return "SY" + s[len(pref):]
    return s


def _fill_record_no(doc, no):
    """把正文「编号：」段落补成「编号：<委托单号>」，保留原字体。"""
    no = _raw_record_no(no)
    if not no:
        return
    for p in doc.paragraphs:
        txt = ''.join(r.text or '' for r in p.runs)
        if txt.strip().startswith('编号'):
            runs = p.runs
            if not runs:
                continue
            # 保留首个 run 的「编号：」标签，把单号接到其后
            base = runs[0].text or ''
            if '：' in base:
                runs[0].text = base.split('：')[0] + '：' + no
            else:
                runs[0].text = base + no
            for r in runs[1:]:
                r.text = ''
            break


def generate_raw_records(project, out_path):
    """project: {info:{...}, tests:[{...}]} -> 为每个测试项生成独立的原始记录文档。
    返回生成的文件路径列表。out_path 作为基准路径，实际文件名会加上测试项目名或序号。"""
    info = project.get("info", {}) or {}
    tests = project.get("tests", []) or []
    if not tests:
        tests = [{}]  # 没有测试项也产出一张空表

    # 解析 out_path: 目录 + 基础文件名
    out_dir = os.path.dirname(out_path)
    base_name = os.path.splitext(os.path.basename(out_path))[0]

    generated_files = []

    for idx, test in enumerate(tests):
        # 为每个测试项生成独立文档
        doc = Document(SKELETON)
        _fill_record_no(doc, info.get("commission_no", ""))
        tbl = doc.tables[0]._tbl
        _fill_table(tbl, info, test, doc)
        _normalize_font_size(doc)  # 全文五号(sz21) 统一压回小五(sz18)

        # 文件命名：基础名_测试项目名.docx 或 基础名_序号.docx
        test_title = test.get("title", "").strip()
        if test_title:
            # 清理文件名非法字符
            safe_title = "".join(c if c.isalnum() or c in (' ', '-', '_', '（', '）', '(', ')') else '_'
                                for c in test_title)
            file_name = f"{base_name}_{safe_title}.docx"
        else:
            file_name = f"{base_name}_{idx+1}.docx"

        file_path = os.path.join(out_dir, file_name)
        os.makedirs(out_dir, exist_ok=True)
        doc.save(file_path)
        generated_files.append(file_path)

    # 返回文件路径列表（保持向后兼容，第一个文件路径也作为主返回值）
    # 始终返回列表，简化调用方逻辑
    return generated_files


def _normalize_font_size(doc):
    """把全文五号(sz=21=10.5pt)统一改成小五(sz=18=9pt)。
    标题(sz30)、编号行(sz24)等其它字号保留不动。
    注意：填入值经共用的 force_song5 设成了 sz21，这里在填表后统一压回小五，
    只影响原始记录，不动主报告。"""
    for r in doc.element.body.iter(W + 'r'):
        rpr = r.find(W + 'rPr')
        if rpr is None:
            continue
        for tag in ('sz', 'szCs'):
            e = rpr.find(W + tag)
            if e is not None and e.get(qn('w:val')) == '21':
                e.set(qn('w:val'), '18')


def _make_page_break(doc):
    """构造一个只含分页符的段落元素。"""
    p = doc.element.body.makeelement(W + 'p', {})
    r = p.makeelement(W + 'r', {})
    br = r.makeelement(W + 'br', {})
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    return p
